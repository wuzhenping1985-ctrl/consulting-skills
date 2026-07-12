#!/usr/bin/env python3
"""Format a DOCX or text-based PDF as a Chinese party/government official document."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt
except ImportError as exc:  # pragma: no cover
    raise SystemExit("Missing dependency: python-docx. Install it or use the bundled Codex document runtime.") from exc


TITLE_FONT = "方正小标宋简体"
TITLE_FALLBACK = "SimSun"
BODY_FONT = "仿宋_GB2312"
BODY_FALLBACK = "FangSong"
SONG_FONT = "SimSun"
HEITI_FONT = "SimHei"
KAITI_FONT = "楷体_GB2312"

DATE_RE = re.compile(r"^[一二三四五六七八九十〇零\d]{4}年[一二三四五六七八九十〇零\d]{1,2}月[一二三四五六七八九十〇零\d]{1,3}日$")
DOC_NO_RE = re.compile(r"[〔\[]\d{4}[〕\]]\s*\d+\s*号")
LEVEL_1_RE = re.compile(r"^[一二三四五六七八九十]+、")
LEVEL_2_RE = re.compile(r"^（[一二三四五六七八九十]+）")
LEVEL_3_RE = re.compile(r"^\d+[.．、]")
CHAPTER_HEADING_RE = re.compile(r"^第[一二三四五六七八九十百千万零〇\d]+[篇章节](?:\s|$|、|：|:)")
CHAPTER_HEADING_SPACE_RE = re.compile(r"^(第\s*[一二三四五六七八九十百千万零〇\d](?:\s*[一二三四五六七八九十百千万零〇\d])*\s*[篇章节])\s+(.+)$")
CJK = r"\u3400-\u4dbf\u4e00-\u9fff"
FULLWIDTH_PUNCT = "，。；：！？、（）《》“”‘’"


def set_run_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_base(paragraph, line_pt: float = 28, first_indent_chars: float = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(line_pt)
    if first_indent_chars:
        fmt.first_line_indent = Pt(16 * first_indent_chars)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("- ")
    set_run_font(run, SONG_FONT, 14)

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = paragraph.add_run()
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    run._r.append(fld_text)

    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run = paragraph.add_run(" -")
    set_run_font(run, SONG_FONT, 14)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(16)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = Pt(28)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_base(paragraph, 14)
    add_page_field(paragraph)


def extract_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    lines: list[str] = []
    for paragraph in doc.paragraphs:
        text = normalize_line(paragraph.text)
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_line(cell.text) for cell in row.cells]
            line = "    ".join(cell for cell in cells if cell)
            if line:
                lines.append(line)
    return lines


def extract_pdf(path: Path) -> list[str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise SystemExit("Missing dependency: pypdf. Convert the PDF to DOCX/text first, or install pypdf.") from exc

    reader = PdfReader(str(path))
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw in text.splitlines():
            line = normalize_line(raw)
            if line:
                lines.append(line)
    if not lines:
        raise SystemExit("No text could be extracted from the PDF. It may be scanned; run OCR first.")
    return lines


def normalize_line(text: str) -> str:
    line = re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()
    return clean_unnecessary_spaces(line)


def clean_unnecessary_spaces(text: str) -> str:
    """Remove accidental Chinese-text spaces while preserving meaningful separators."""
    chapter_match = CHAPTER_HEADING_SPACE_RE.match(text)
    if chapter_match:
        prefix = re.sub(r"\s+", "", chapter_match.group(1))
        suffix = clean_unnecessary_spaces(chapter_match.group(2))
        return f"{prefix} {suffix}" if suffix else prefix

    text = re.sub(fr"(?<=[{CJK}])\s+(?=[{CJK}])", "", text)
    text = re.sub(fr"(?<=[{CJK}\d])\s+(?=[{FULLWIDTH_PUNCT}])", "", text)
    text = re.sub(fr"(?<=[{FULLWIDTH_PUNCT}])\s+(?=[{CJK}\d])", "", text)
    text = re.sub(fr"(?<=[{CJK}])\s+(?=\d)", "", text)
    text = re.sub(fr"(?<=\d)\s+(?=[{CJK}])", "", text)
    return text.strip()


def classify_line(line: str, index: int, title_index: int | None) -> str:
    if title_index is not None and index == title_index:
        return "title"
    if DOC_NO_RE.search(line) and len(line) <= 40:
        return "doc_no"
    if line.startswith("附件：") or line.startswith("附件:"):
        return "attachment_note"
    if DATE_RE.match(line):
        return "date"
    if CHAPTER_HEADING_RE.match(line):
        return "chapter_heading"
    if LEVEL_1_RE.match(line):
        return "heading_1"
    if LEVEL_2_RE.match(line):
        return "heading_2"
    if LEVEL_3_RE.match(line):
        return "heading_3"
    if looks_like_signature(line):
        return "signature"
    return "body"


def looks_like_signature(line: str) -> bool:
    if len(line) > 35:
        return False
    suffixes = ("委员会", "人民政府", "办公厅", "办公室", "厅", "局", "部", "委", "院", "会", "公司")
    return line.endswith(suffixes) and not line.endswith(("如下", "工作", "问题"))


def choose_title_index(lines: list[str], explicit_title: str | None) -> int | None:
    if explicit_title:
        return None
    for i, line in enumerate(lines[:12]):
        if DOC_NO_RE.search(line):
            continue
        if line.endswith(("号", "日")) and len(line) <= 24:
            continue
        if len(line) >= 4:
            return i
    return None


def add_text_paragraph(document: Document, text: str, kind: str) -> None:
    paragraph = document.add_paragraph()

    if kind == "blank":
        set_paragraph_base(paragraph, 28)
        return

    if kind == "title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_base(paragraph, 28)
        run = paragraph.add_run(text)
        set_run_font(run, TITLE_FONT, 22, bold=False)
        return

    if kind == "doc_no":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_base(paragraph, 28)
        run = paragraph.add_run(text)
        set_run_font(run, BODY_FONT, 16)
        return

    if kind in {"signature", "date"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_base(paragraph, 28)
        run = paragraph.add_run(text)
        set_run_font(run, BODY_FONT, 16)
        return

    if kind == "chapter_heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_base(paragraph, 28)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font(run, HEITI_FONT, 16)
        return

    if kind == "heading_1":
        set_paragraph_base(paragraph, 28, first_indent_chars=2)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font(run, HEITI_FONT, 16)
        return

    if kind == "heading_2":
        set_paragraph_base(paragraph, 28, first_indent_chars=2)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font(run, KAITI_FONT, 16)
        return

    if kind == "heading_3":
        set_paragraph_base(paragraph, 28, first_indent_chars=2)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(text)
        set_run_font(run, BODY_FONT, 16, bold=True)
        return

    first_indent = 0 if kind == "attachment_note" else 2
    set_paragraph_base(paragraph, 28, first_indent_chars=first_indent)
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, 16)


def add_blank_paragraph_if_needed(document: Document) -> None:
    if not document.paragraphs or document.paragraphs[-1].text.strip():
        add_text_paragraph(document, "", "blank")


def build_document(lines: list[str], output: Path, title: str | None = None) -> None:
    document = Document()
    configure_document(document)

    if title:
        add_text_paragraph(document, title, "title")
        add_blank_paragraph_if_needed(document)

    title_index = choose_title_index(lines, title)
    for i, line in enumerate(lines):
        kind = classify_line(line, i, title_index)
        if kind == "chapter_heading":
            add_blank_paragraph_if_needed(document)
        add_text_paragraph(document, line, kind)
        if kind in {"title", "chapter_heading"}:
            add_blank_paragraph_if_needed(document)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Format DOCX/PDF text as GB/T 9704-2012-style official DOCX.")
    parser.add_argument("-i", "--input", required=True, type=Path, help="Input .docx or text-based .pdf file")
    parser.add_argument("-o", "--output", required=True, type=Path, help="Output .docx path")
    parser.add_argument("--title", help="Override or provide the main document title")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    suffix = args.input.suffix.lower()
    if suffix == ".docx":
        lines = extract_docx(args.input)
    elif suffix == ".pdf":
        lines = extract_pdf(args.input)
    elif suffix == ".doc":
        raise SystemExit("Legacy .doc input is not handled directly. Convert it to .docx first, then rerun.")
    else:
        raise SystemExit(f"Unsupported input type: {suffix}. Use .docx or text-based .pdf.")

    if not lines and not args.title:
        raise SystemExit("No text found in input.")
    build_document(lines, args.output, args.title)
    print(f"Saved formatted DOCX: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
