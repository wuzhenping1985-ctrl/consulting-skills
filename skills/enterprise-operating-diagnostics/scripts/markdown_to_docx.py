#!/usr/bin/env python
"""Convert a simple Chinese Markdown report to a formatted DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, line_spacing: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_markdown_runs(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(width):
            cell = table.cell(i, j)
            text = row[j].strip() if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_format(p, line_spacing=1.1)
            run = p.add_run(text)
            set_run_font(run, size=10.5, bold=(i == 0))


def flush_table(document: Document, table_rows: list[list[str]]) -> list[list[str]]:
    if table_rows:
        add_table(document, table_rows)
    return []


def parse_table_row(line: str) -> list[str] | None:
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return None
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
        return []
    return cells


def convert(markdown_path: Path, docx_path: Path) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    section = document.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    table_rows: list[list[str]] = []
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    for line in lines:
        row = parse_table_row(line)
        if row is not None:
            if row:
                table_rows.append(row)
            continue

        table_rows = flush_table(document, table_rows)
        text = line.strip()
        if not text:
            continue

        if text.startswith("# "):
            p = document.add_paragraph()
            p.style = styles["Title"]
            set_paragraph_format(p, line_spacing=1.15)
            run = p.add_run(text[2:].strip())
            set_run_font(run, size=18, bold=True)
        elif text.startswith("## "):
            p = document.add_heading(level=1)
            set_paragraph_format(p, line_spacing=1.2)
            run = p.add_run(text[3:].strip())
            set_run_font(run, size=15, bold=True)
        elif text.startswith("### "):
            p = document.add_heading(level=2)
            set_paragraph_format(p, line_spacing=1.2)
            run = p.add_run(text[4:].strip())
            set_run_font(run, size=13, bold=True)
        elif text.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            set_paragraph_format(p)
            add_markdown_runs(p, text[2:].strip(), size=11)
        elif re.match(r"^\d+[.、]\s", text):
            p = document.add_paragraph(style="List Number")
            set_paragraph_format(p)
            add_markdown_runs(p, re.sub(r"^\d+[.、]\s*", "", text), size=11)
        else:
            p = document.add_paragraph()
            set_paragraph_format(p)
            add_markdown_runs(p, text, size=11)

    flush_table(document, table_rows)

    for paragraph in document.paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        jc = p_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            p_pr.append(jc)
        if paragraph.style.name == "Normal":
            jc.set(qn("w:val"), "both")

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX for enterprise diagnostic reports.")
    parser.add_argument("input", type=Path, help="UTF-8 Markdown input file")
    parser.add_argument("output", type=Path, help="DOCX output file")
    args = parser.parse_args()
    convert(args.input, args.output)


if __name__ == "__main__":
    main()
